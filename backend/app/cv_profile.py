"""Secure CV ingestion and conservative profile extraction."""
from __future__ import annotations

from dataclasses import dataclass
from io import BytesIO
from pathlib import PurePosixPath
from typing import Any, BinaryIO
import re
import zipfile

MAX_CV_BYTES = 10 * 1024 * 1024
SUPPORTED_EXTENSIONS = {".pdf", ".docx"}
SUPPORTED_MIME_TYPES = {
    ".pdf": {"application/pdf", "application/octet-stream"},
    ".docx": {"application/vnd.openxmlformats-officedocument.wordprocessingml.document", "application/octet-stream"},
}
_SECTIONS = {
    "skills": "skills", "technical skills": "skills", "habilidades": "skills",
    "experience": "experience", "professional experience": "experience", "work experience": "experience", "experiencia": "experience",
    "education": "education", "formación": "education", "formacion": "education", "educacion": "education",
    "languages": "languages", "idiomas": "languages",
}
_KNOWN_SKILLS = ("python", "javascript", "typescript", "java", "c#", ".net", "go", "rust", "sql", "postgresql", "mysql", "fastapi", "django", "react", "next.js", "node.js", "docker", "kubernetes", "aws", "azure", "gcp", "git")


class CVValidationError(ValueError):
    """The uploaded CV is unsupported, unsafe, malformed, or unreadable."""


class CVParserUnavailable(RuntimeError):
    """The optional parser dependency for the selected format is missing."""


@dataclass(frozen=True, slots=True)
class ParsedCV:
    filename: str
    content_type: str | None
    text: str
    profile: dict[str, Any]


def _validate_docx(data: bytes, max_bytes: int) -> None:
    try:
        with zipfile.ZipFile(BytesIO(data)) as archive:
            names = archive.namelist()
            if "word/document.xml" not in names:
                raise CVValidationError("DOCX is missing word/document.xml")
            total = 0
            for info in archive.infolist():
                path = PurePosixPath(info.filename)
                if info.flag_bits & 1 or path.is_absolute() or ".." in path.parts:
                    raise CVValidationError("DOCX contains an unsafe or encrypted ZIP member")
                total += info.file_size
                if total > max_bytes * 20:
                    raise CVValidationError("DOCX contains excessive uncompressed data")
    except zipfile.BadZipFile as exc:
        raise CVValidationError("CV is not a valid DOCX file") from exc


def _validate(data: bytes, filename: str, content_type: str | None, max_bytes: int) -> str:
    if not filename or len(filename) > 255 or filename != PurePosixPath(filename).name or "\\" in filename:
        raise CVValidationError("CV filename must be a safe basename of at most 255 characters")
    if any(ord(char) < 32 for char in filename):
        raise CVValidationError("CV filename contains control characters")
    extension = PurePosixPath(filename).suffix.lower()
    if extension not in SUPPORTED_EXTENSIONS:
        raise CVValidationError("Only PDF and DOCX CV files are supported")
    if not data:
        raise CVValidationError("CV file is empty")
    if len(data) > max_bytes:
        raise CVValidationError(f"CV file exceeds the {max_bytes} byte limit")
    mime = (content_type or "").split(";", 1)[0].strip().lower()
    if mime and mime not in SUPPORTED_MIME_TYPES[extension]:
        raise CVValidationError("CV MIME type does not match its extension")
    if extension == ".pdf" and not data.startswith(b"%PDF-"):
        raise CVValidationError("CV is not a valid PDF file")
    if extension == ".docx":
        _validate_docx(data, max_bytes)
    return filename


def _extract(data: bytes, extension: str) -> str:
    if extension == ".pdf":
        try:
            from pypdf import PdfReader
        except ImportError as exc:
            raise CVParserUnavailable("Install pypdf to process PDF CV files") from exc
        try:
            reader = PdfReader(BytesIO(data), strict=False)
            if reader.is_encrypted:
                raise CVValidationError("Encrypted PDFs are not accepted")
            return "\n".join(page.extract_text() or "" for page in reader.pages)
        except CVValidationError:
            raise
        except Exception as exc:
            raise CVValidationError("PDF could not be read") from exc
    try:
        from docx import Document
    except ImportError as exc:
        raise CVParserUnavailable("Install python-docx to process DOCX CV files") from exc
    try:
        document = Document(BytesIO(data))
        lines = [paragraph.text for paragraph in document.paragraphs]
        lines.extend(cell.text for table in document.tables for row in table.rows for cell in row.cells)
        return "\n".join(lines)
    except Exception as exc:
        raise CVValidationError("DOCX could not be read") from exc


def _section(line: str) -> str | None:
    return _SECTIONS.get(re.sub(r"[^a-záéíóúñ ]", "", line.lower()).strip())


def _profile(text: str) -> dict[str, Any]:
    lines = [line.strip() for line in text.splitlines() if line.strip()]
    name = next((line for line in lines[:5] if len(line) <= 100 and not _section(line)), "")
    sections: dict[str, list[str]] = {"summary": []}
    current = "summary"
    for line in lines:
        key = _section(line)
        if key:
            current = key
            sections.setdefault(key, [])
        else:
            sections.setdefault(current, []).append(line)
    source = " ".join(sections.get("skills", [])) or text
    skills = [skill for skill in _KNOWN_SKILLS if skill.lower() in source.lower()]
    if sections.get("skills"):
        for skill_piece in re.split(r"[,;|•]", " ".join(sections["skills"])):
            skill_piece = skill_piece.strip()
            if skill_piece and len(skill_piece) <= 60 and skill_piece.lower() not in {skill.lower() for skill in skills}:
                skills.append(skill_piece)
    return {
        "name": name,
        "skills": skills,
        "experience": [{"text": value} for value in sections.get("experience", [])],
        "education": [{"text": value} for value in sections.get("education", [])],
        "languages": sections.get("languages", []),
        "summary": " ".join(sections["summary"][:3]),
    }


def parse_cv(file: BinaryIO, filename: str, content_type: str | None = None, *, max_bytes: int = MAX_CV_BYTES) -> ParsedCV:
    """Validate, extract, and convert a PDF/DOCX stream to an editable profile."""
    data = file.read(max_bytes + 1)
    safe_name = _validate(data, filename, content_type, max_bytes)
    text = _extract(data, PurePosixPath(safe_name).suffix.lower())
    text = "\n".join(re.sub(r"[ \t]+", " ", line).strip() for line in text.replace("\x00", "").splitlines()).strip()
    text = re.sub(r"\n{3,}", "\n\n", text)
    if len(text) < 20:
        raise CVValidationError("CV text is empty or unreadable")
    return ParsedCV(safe_name, content_type, text, _profile(text))
